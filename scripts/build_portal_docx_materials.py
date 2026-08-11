from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
PORTAL = ROOT / "submission" / "nature_communications" / "portal_materials_final"


def set_run_font(run, name="Arial", size=11, bold=False, italic=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_para_spacing(paragraph, before=0, after=8, line=1.15):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_hyperlink_style(document):
    styles = document.styles
    if "Hyperlink" not in styles:
        style = styles.add_style("Hyperlink", 2)
        style.font.color.rgb = RGBColor(5, 99, 193)
        style.font.underline = True


def add_field_page_number(paragraph):
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def configure_doc(document):
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    for style_name, size, color, before, after in [
        ("Heading 1", 20, "000000", 20, 6),
        ("Heading 2", 16, "000000", 18, 6),
        ("Heading 3", 14, "434343", 16, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = False
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.15

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_para_spacing(footer, 0, 0, 1.0)
    run = footer.add_run("Page ")
    set_run_font(run, size=9, color="555555")
    add_field_page_number(footer)
    add_hyperlink_style(document)


def add_formatted_runs(paragraph, text):
    # Keep Markdown backticks visible as plain code markers only where they denote paths.
    pieces = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)", text)
    for piece in pieces:
        if not piece:
            continue
        if piece.startswith("**") and piece.endswith("**"):
            run = paragraph.add_run(piece[2:-2])
            set_run_font(run, bold=True)
        elif piece.startswith("*") and piece.endswith("*"):
            run = paragraph.add_run(piece[1:-1])
            set_run_font(run, italic=True)
        elif piece.startswith("`") and piece.endswith("`"):
            run = paragraph.add_run(piece[1:-1])
            set_run_font(run, name="Courier New", size=10)
        else:
            run = paragraph.add_run(piece)
            set_run_font(run)


def markdown_to_docx(md_path, out_path, subtitle=None):
    document = Document()
    configure_doc(document)

    lines = md_path.read_text(encoding="utf-8").splitlines()
    in_code = False
    code_lines = []
    first_heading_done = False

    for raw in lines:
        line = raw.rstrip()

        if line.startswith("```"):
            if in_code:
                if code_lines:
                    paragraph = document.add_paragraph()
                    paragraph.style = document.styles["Normal"]
                    set_para_spacing(paragraph, 4, 8, 1.0)
                    run = paragraph.add_run("\n".join(code_lines))
                    set_run_font(run, name="Courier New", size=9, color="333333")
                code_lines = []
                in_code = False
            else:
                in_code = True
            continue

        if in_code:
            code_lines.append(line)
            continue

        if not line.strip():
            continue

        if line.startswith("# "):
            text = line[2:].strip()
            if not first_heading_done:
                paragraph = document.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                set_para_spacing(paragraph, 0, 3, 1.15)
                run = paragraph.add_run(text)
                set_run_font(run, size=26, bold=False, color="000000")
                first_heading_done = True
                if subtitle:
                    sub = document.add_paragraph()
                    set_para_spacing(sub, 0, 14, 1.15)
                    run = sub.add_run(subtitle)
                    set_run_font(run, size=11, color="555555")
            else:
                document.add_heading(text, level=1)
            continue

        if line.startswith("## "):
            document.add_heading(line[3:].strip(), level=2)
            continue

        if line.startswith("### "):
            document.add_heading(line[4:].strip(), level=3)
            continue

        if line.startswith("- "):
            paragraph = document.add_paragraph(style="List Bullet")
            set_para_spacing(paragraph, 0, 4, 1.15)
            add_formatted_runs(paragraph, line[2:].strip())
            continue

        m = re.match(r"^(\d+)\.\s+(.*)$", line)
        if m:
            paragraph = document.add_paragraph(style="List Number")
            set_para_spacing(paragraph, 0, 4, 1.15)
            add_formatted_runs(paragraph, m.group(2).strip())
            continue

        if line.startswith("> "):
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.25)
            set_para_spacing(paragraph, 4, 8, 1.15)
            run = paragraph.add_run(line[2:].strip())
            set_run_font(run, italic=True, color="555555")
            continue

        paragraph = document.add_paragraph()
        set_para_spacing(paragraph, 0, 8, 1.15)
        add_formatted_runs(paragraph, line)

    document.save(out_path)


def build_cover_letter(md_path, out_path):
    document = Document()
    configure_doc(document)

    title = document.add_paragraph()
    set_para_spacing(title, 0, 16, 1.15)
    run = title.add_run("Cover Letter")
    set_run_font(run, size=26, color="000000")

    lines = md_path.read_text(encoding="utf-8").splitlines()
    for line in lines:
        if not line.strip():
            continue
        paragraph = document.add_paragraph()
        set_para_spacing(paragraph, 0, 8, 1.15)
        add_formatted_runs(paragraph, line)

    document.save(out_path)


def main():
    markdown_to_docx(
        PORTAL / "Supplementary_Information_V3.md",
        PORTAL / "Supplementary_Information_V3.docx",
        subtitle="Supplementary material for Nature Communications submission",
    )
    build_cover_letter(
        PORTAL / "COVER_LETTER_V8_FINAL.md",
        PORTAL / "COVER_LETTER_V8_FINAL.docx",
    )


if __name__ == "__main__":
    main()
