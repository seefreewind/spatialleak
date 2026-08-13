from pathlib import Path
import json

import matplotlib as mpl
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SUB = ROOT / "submission" / "nature_communications"
PORTAL = SUB / "portal_materials_final"
SD = PORTAL / "source_data"
FIGDIR = PORTAL / "supplementary_figures"
FIGDIR.mkdir(parents=True, exist_ok=True)


mpl.rcParams.update({
    "font.family": "sans-serif",
    "font.sans-serif": ["Arial", "Helvetica", "DejaVu Sans", "sans-serif"],
    "pdf.fonttype": 42,
    "svg.fonttype": "none",
    "font.size": 7,
    "axes.spines.right": False,
    "axes.spines.top": False,
    "axes.linewidth": 0.7,
    "legend.frameon": False,
})


COLORS = {
    "random": "#4C566A",
    "spatial": "#5E81AC",
    "subject": "#BF616A",
    "size": "#A3BE8C",
    "buffer": "#D08770",
    "pca": "#5E81AC",
    "knn": "#88C0D0",
    "sage": "#B48EAD",
}


def round_float(x, digits=3):
    if x is None:
        return ""
    try:
        if pd.isna(x):
            return ""
    except TypeError:
        pass
    try:
        return f"{float(x):.{digits}f}"
    except Exception:
        return str(x)


def save_fig(fig, stem):
    out = FIGDIR / stem
    fig.savefig(out.with_suffix(".pdf"), bbox_inches="tight")
    fig.savefig(out.with_suffix(".png"), dpi=300, bbox_inches="tight")
    plt.close(fig)
    return out.with_suffix(".png")


def build_supp_fig1():
    df = pd.read_csv(SD / "SupplementaryFigure1_SourceData.csv")
    out_csv = SD / "SupplementaryFigure1_SourceData.csv"
    labels = [f"{r.model}\n{r.dataset}" for r in df.itertuples()]
    x = np.arange(len(df))
    width = 0.25
    fig, ax = plt.subplots(figsize=(7.2, 3.6))
    ax.bar(x - width, df["random"], width, label="Random", color=COLORS["random"])
    ax.bar(x, df["spatial_strict"], width, label="Spatial strict", color=COLORS["spatial"])
    subject_vals = df["patient_strict"].replace("", np.nan).astype(float)
    ax.bar(x + width, subject_vals, width, label="Subject-held-out", color=COLORS["subject"])
    ax.set_ylabel("Mean Pearson correlation")
    ax.set_xticks(x)
    ax.set_xticklabels(labels, rotation=35, ha="right")
    ax.set_ylim(bottom=min(0, np.nanmin(df[["random", "spatial_strict"]].to_numpy()) - 0.04))
    ax.legend(ncol=3, loc="upper right")
    ax.axhline(0, color="#333333", lw=0.6)
    ax.set_title("Evaluation regime changes apparent model performance", loc="left", fontsize=9)
    fig.tight_layout()
    return save_fig(fig, "Supplementary_Fig1_evaluation_regime_behavior"), out_csv


def build_supp_fig2():
    df = pd.read_csv(ROOT / "results" / "paper_assets" / "table_random_size_matched_control.csv")
    keep = df[df["hop"].eq(5)].copy()
    agg = keep.groupby(["dataset_label", "model"], as_index=False).agg(
        random_full_mean=("random_full_mean_pearson", "mean"),
        random_size_matched_mean=("random_size_matched_mean_pearson", "mean"),
        spatial_buffer_mean=("spatial_buffer_mean_pearson", "mean"),
        delta_size=("delta_size", "mean"),
        delta_spatial=("delta_spatial", "mean"),
    )
    agg.to_csv(SD / "SupplementaryFigure2_SourceData.csv", index=False)

    agg["label"] = agg["dataset_label"] + " | " + agg["model"].replace({"pca_ridge": "PCA+Ridge", "spatial_knn": "Spatial kNN"})
    x = np.arange(len(agg))
    fig, ax = plt.subplots(figsize=(7.2, 3.4))
    width = 0.34
    ax.bar(x - width / 2, agg["delta_size"], width, color=COLORS["size"], label="Random size-matched loss")
    ax.bar(x + width / 2, agg["delta_spatial"], width, color=COLORS["buffer"], label="Spatial-buffer loss")
    ax.set_ylabel("Mean Pearson loss")
    ax.set_xticks(x)
    ax.set_xticklabels(agg["label"], rotation=35, ha="right")
    ax.axhline(0, color="#333333", lw=0.6)
    ax.set_title("Spatial-buffer losses exceed sample-count effects", loc="left", fontsize=9)
    ax.legend(loc="upper left")
    fig.tight_layout()
    return save_fig(fig, "Supplementary_Fig2_sample_size_matched"), SD / "SupplementaryFigure2_SourceData.csv"


def build_supp_fig3():
    base = pd.read_csv(ROOT / "results" / "paper_assets" / "table_shared_panel50_RLI.csv")
    sage = pd.read_csv(ROOT / "results" / "paper_assets" / "table_graphsage_shared_panel50_RLI_trainonly.csv")
    base = base[base["model"].isin(["pca_ridge", "spatial_knn"]) & base["RLI"].notna()].copy()
    base["preprocessing"] = "fixed_panel"
    sage = sage.rename(columns={"strict_label": "strict_label"})
    combined = pd.concat([
        base[["dataset", "strict_label", "model", "random", "strict", "LI", "RLI", "retention", "preprocessing"]],
        sage[["dataset", "strict_label", "model", "random", "strict", "LI", "RLI", "retention", "preprocessing"]],
    ], ignore_index=True)
    combined = combined[combined["RLI"].notna()].copy()
    combined.to_csv(SD / "SupplementaryFigure3_SourceData.csv", index=False)

    plot = combined.copy()
    plot["label"] = plot["dataset"] + "\n" + plot["model"].replace({
        "pca_ridge": "PCA+Ridge", "spatial_knn": "Spatial kNN", "graphsage": "GraphSAGE"
    }) + " | " + plot["strict_label"]
    plot = plot.sort_values(["dataset", "model", "strict_label"])
    x = np.arange(len(plot))
    colors = plot["model"].map({"pca_ridge": COLORS["pca"], "spatial_knn": COLORS["knn"], "graphsage": COLORS["sage"]}).fillna("#777777")
    fig, ax = plt.subplots(figsize=(8.0, 3.8))
    ax.bar(x, plot["RLI"], color=colors)
    ax.axhline(0, color="#333333", lw=0.6)
    ax.axhline(0.5, color="#999999", lw=0.5, ls="--")
    ax.set_ylabel("Relative leakage inflation")
    ax.set_xticks(x)
    ax.set_xticklabels(plot["label"], rotation=55, ha="right")
    ax.set_title("shared_panel_50 reproduces subject- and spatial-channel patterns", loc="left", fontsize=9)
    fig.tight_layout()
    return save_fig(fig, "Supplementary_Fig3_shared_panel50"), SD / "SupplementaryFigure3_SourceData.csv"


def make_tables():
    two = pd.read_csv(ROOT / "results" / "paper_assets" / "table_two_channel_leakage.csv")
    datasets = []
    for dataset in ["DLPFC", "Andersson", "Thrane", "Visium breast"]:
        sub = two[two["dataset"].eq(dataset)].iloc[0]
        datasets.append({
            "Dataset": dataset,
            "Platform": sub["platform"],
            "Patients/donors": int(sub["patient_count"]),
            "Sections": int(sub["section_count"]),
            "Role in manuscript": {
                "DLPFC": "primary spatial and donor/section evaluation",
                "Andersson": "subject-associated breast cancer evaluation",
                "Thrane": "subject-associated melanoma evaluation",
                "Visium breast": "dense Visium spatial/section evaluation",
            }[dataset],
        })
    gse = pd.read_csv(ROOT / "data" / "external_audit" / "gse278936" / "public_sample_audit.csv")
    datasets.append({
        "Dataset": "GSE278936 prostate",
        "Platform": "10x Visium",
        "Patients/donors": gse["patient_id"].nunique(),
        "Sections": gse["section_id"].nunique(),
        "Role in manuscript": "external spatial-channel replication only",
    })
    t1 = pd.DataFrame(datasets)
    t1.to_csv(SD / "SupplementaryTable1_DatasetSampleStructure.csv", index=False)

    splits = pd.read_csv(ROOT / "results" / "paper_assets" / "table_split_sample_sizes.csv")
    t2 = splits.groupby(["dataset_label", "split"], as_index=False).agg(
        n_total=("n_total", "first"),
        mean_train=("n_train", "mean"),
        mean_validation=("n_val", "mean"),
        mean_test=("n_test", "mean"),
        mean_dropped=("n_dropped", "mean"),
        target_genes=("target_genes", "first"),
    )
    for c in ["mean_train", "mean_validation", "mean_test", "mean_dropped"]:
        t2[c] = t2[c].round(1)
    t2.to_csv(SD / "SupplementaryTable2_SplitSampleCounts.csv", index=False)

    t3 = pd.read_csv(SD / "SupplementaryTable_VisiumBreastSectionHeldOut.csv")
    t3sum = t3.groupby(["split", "model"], as_index=False).agg(
        mean_pearson=("mean_pearson", "mean"),
        median_pearson=("median_pearson", "mean"),
        mean_spearman=("mean_spearman", "mean"),
        mean_rmse=("mean_rmse", "mean"),
        n_genes=("n_genes", "first"),
    )
    for c in ["mean_pearson", "median_pearson", "mean_spearman", "mean_rmse"]:
        t3sum[c] = t3sum[c].round(3)
    t3sum.to_csv(SD / "SupplementaryTable3_VisiumBreastSectionHeldOutSummary.csv", index=False)

    mixed = json.loads((ROOT / "results" / "final_stats" / "mixed_effects.json").read_text())
    rows = []
    for channel in ["patient", "spatial"]:
        d = mixed[channel]
        rows.append({
            "Channel": "subject-associated" if channel == "patient" else "spatial",
            "n observations": d["n_obs"],
            "n datasets": d["n_datasets"],
            "Moran I coefficient": round(d["coef_moran_i"], 3),
            "Moran I P value": f"{d['p_moran_i']:.3g}",
            "Model fixed effect shown": "Spatial kNN vs PCA+Ridge",
            "Model coefficient": round(d["coefs"].get("C(model)[T.spatial_knn]", float("nan")), 3),
            "Model P value": f"{d['pvalues'].get('C(model)[T.spatial_knn]', float('nan')):.3g}",
        })
    pd.DataFrame(rows).to_csv(SD / "SupplementaryTable4_MixedEffectsSummary.csv", index=False)

    wilcoxon = pd.read_csv(ROOT / "results" / "final_stats" / "wilcoxon_all_datasets.csv")
    rli = pd.read_csv(ROOT / "results" / "final_stats" / "LI_RLI_all_datasets.csv")
    rli = rli[["dataset", "strict_type", "model", "random", "strict", "LI", "RLI", "retention"]]
    t6 = wilcoxon.merge(rli, on=["dataset", "strict_type", "model"], how="left")
    t6 = t6.rename(columns={
        "dataset": "Dataset",
        "strict_type": "Comparison family",
        "model": "Model",
        "n_pairs": "Paired n",
        "median_diff": "Median random-minus-strict difference",
        "statistic": "Wilcoxon W",
        "p": "P value",
        "p_bh": "BH-FDR q value",
        "sig_bh_0.05": "BH-FDR q < 0.05",
    })
    t6.insert(3, "Test", "Two-sided paired Wilcoxon signed-rank test")
    t6.insert(4, "Null hypothesis", "Median paired random-minus-strict difference = 0")
    t6.insert(7, "Degrees of freedom", "n/a")
    low_random = t6["random"].abs() < 0.05
    t6.loc[low_random, ["RLI", "retention"]] = np.nan
    t6["Effect metric"] = "median paired difference; LI/RLI/retention, with RLI and retention left blank when absolute random mean Pearson < 0.05"
    for col in ["Median random-minus-strict difference", "Wilcoxon W", "P value", "BH-FDR q value", "random", "strict", "LI", "RLI", "retention"]:
        t6[col] = t6[col].map(lambda x: "" if pd.isna(x) else f"{float(x):.6g}")
    t6["BH-FDR q < 0.05"] = t6["BH-FDR q < 0.05"].map({True: "yes", False: "no"})
    t6.to_csv(SD / "SupplementaryTable6_WilcoxonSignedRankResults.csv", index=False)

    fig3 = pd.read_csv(SD / "Figure3_SourceData.csv")
    gse_rli = pd.read_csv(ROOT / "results" / "paper_assets" / "table_gse278936_spatial_pilot_RLI.csv")
    rows = []
    for r in fig3.itertuples():
        if isinstance(r.spatial_na_reason, str) and r.spatial_na_reason:
            rows.append({
                "Dataset": r.dataset,
                "Model": r.model,
                "Channel/comparison": "spatial",
                "Reason": r.spatial_na_reason,
            })
        if isinstance(r.patient_na_reason, str) and r.patient_na_reason:
            rows.append({
                "Dataset": r.dataset,
                "Model": r.model,
                "Channel/comparison": "subject-associated",
                "Reason": r.patient_na_reason,
            })
    for r in gse_rli[gse_rli["rli"].isna()].itertuples():
        rows.append({
            "Dataset": r.dataset,
            "Model": r.model,
            "Channel/comparison": r.comparison,
            "Reason": r.decision_note,
        })
    t5 = pd.DataFrame(rows).drop_duplicates()
    t5.to_csv(SD / "SupplementaryTable5_BoundaryCases.csv", index=False)
    return t1, t2, t3sum, pd.DataFrame(rows).drop_duplicates(), t6


def set_run_font(run, name="Arial", size=10, bold=False, italic=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_para_spacing(paragraph, before=0, after=6, line=1.1):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def configure_doc(document):
    section = document.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for style_name, size, color, before, after in [
        ("Heading 1", 18, "000000", 16, 5),
        ("Heading 2", 14, "000000", 12, 4),
        ("Heading 3", 11, "434343", 8, 3),
    ]:
        style = styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = False
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_text(document, text, bold=False, italic=False, size=10):
    p = document.add_paragraph()
    set_para_spacing(p)
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold, italic=italic)
    return p


def add_caption(document, title, caption, provenance=None):
    p = document.add_paragraph()
    set_para_spacing(p, before=4, after=3, line=1.05)
    r = p.add_run(title + " ")
    set_run_font(r, size=9, bold=True)
    r = p.add_run(caption)
    set_run_font(r, size=9)
    if provenance:
        r = p.add_run(" Source data: " + provenance + ".")
        set_run_font(r, size=8, italic=True, color="555555")


def add_df_table(document, df, max_rows=None, font_size=7.5):
    if max_rows is not None:
        df = df.head(max_rows)
    table = document.add_table(rows=1, cols=len(df.columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, col in enumerate(df.columns):
        cell = table.rows[0].cells[i]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(str(col))
        set_run_font(r, size=font_size, bold=True)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(df.columns):
            val = row[col]
            if isinstance(val, (float, np.floating)):
                if str(col).startswith("mean_") and col in {"mean_train", "mean_validation", "mean_test", "mean_dropped"}:
                    val = round_float(val, 1)
                else:
                    val = round_float(val, 3)
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_run_font(r, size=font_size)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    return table


def build_si_doc(fig_paths, tables):
    t1, t2, t3, t5_dedup, t6 = tables
    t4 = pd.read_csv(SD / "SupplementaryTable4_MixedEffectsSummary.csv")
    doc = Document()
    configure_doc(doc)
    title = doc.add_paragraph()
    set_para_spacing(title, 0, 3, 1.1)
    r = title.add_run("Supplementary Information")
    set_run_font(r, size=24)
    sub = doc.add_paragraph()
    set_para_spacing(sub, 0, 10, 1.1)
    r = sub.add_run("SpatialLeak: evaluation design reshapes apparent generalization in spatial omics prediction")
    set_run_font(r, size=13, color="555555")

    doc.add_heading("Supplementary Methods", level=1)
    for paragraph in [
        "SpatialLeak used public DLPFC, Andersson HER2-positive breast cancer, Thrane melanoma, two public 10x Genomics Visium breast cancer sections and GSE278936 prostate Visium data. Restricted EGA data from the prostate study were not used. All datasets were normalized with library-size scaling to 10,000 counts per spot followed by log1p transformation. Target panels were fixed after normalization. Up to 2,000 highly variable genes (HVGs) were then selected once during dataset preprocessing after excluding target genes and frozen before split-specific model fitting.",
        "Random splits used 80/10/10 train/validation/test proportions. Matched spatial splits used 3 x 3 within-slide grid blocks and 300 candidate assignments per seed. Block-only splits correspond to hop0. Hop2 and hop5 retained test spots only when the nearest-training shortest-path distance was at least two or five edges, respectively, on a within-slide spatial kNN graph with k = 15.",
        "PCA+Ridge used up to 2,000 predictor genes, 64 PCs and Ridge alpha 1.0, with PCA fitted on training observations only. Spatial kNN used k = 15 training neighbours and inverse-distance weighting in normalized per-slide coordinates. GraphSAGE used train-only PCA and scaling, two layers, hidden dimension 128, graph k = 10 with self-loops, ReLU activation, no dropout, mean-squared-error loss on training nodes, Adam learning rate 10^-3, weight decay 10^-4, 500 maximum epochs and validation-loss early stopping with patience 60.",
        "Main DLPFC, Andersson, Thrane and Visium breast baseline analyses used seeds 0-9; GSE278936 spatial-channel replication used seeds 0-4. RLI was not interpreted when absolute random mean Pearson was below 0.05.",
    ]:
        add_text(doc, paragraph)

    doc.add_heading("Supplementary Figures", level=1)
    for idx, (img, src, title, caption) in enumerate([
        (fig_paths[0], "source_data/SupplementaryFigure1_SourceData.csv", "Supplementary Fig. 1 | Evaluation-regime-dependent model behaviour.", "Mean Pearson correlation is shown for random, spatial strict and subject-held-out regimes where defined. Missing subject-held-out bars indicate dataset structures without an interpretable subject-held-out comparison."),
        (fig_paths[1], "source_data/SupplementaryFigure2_SourceData.csv", "Supplementary Fig. 2 | Random-size-matched controls.", "Loss caused by downsampling random splits is compared with the matched spatial-buffer loss at hop5. The spatial-buffer losses remain larger than the sample-count losses in the main spatial-channel settings."),
        (fig_paths[2], "source_data/SupplementaryFigure3_SourceData.csv", "Supplementary Fig. 3 | shared_panel_50 robustness.", "Relative leakage inflation (RLI) is shown for frozen shared_panel_50 analyses, including train-only GraphSAGE preprocessing where applicable."),
    ]):
        if idx:
            doc.add_page_break()
        doc.add_picture(str(img), width=Inches(6.15))
        add_caption(doc, title, caption, provenance=src)

    doc.add_page_break()
    doc.add_heading("Supplementary Tables", level=1)
    add_caption(doc, "Supplementary Table 1 | Dataset and sample structure.", "Datasets used in the locked manuscript evidence hierarchy.", "source_data/SupplementaryTable1_DatasetSampleStructure.csv")
    add_df_table(doc, t1, font_size=7.2)

    doc.add_paragraph()
    add_caption(doc, "Supplementary Table 2 | Split sample counts.", "Mean split counts across frozen seeds for random and spatial-buffer splits.", "source_data/SupplementaryTable2_SplitSampleCounts.csv")
    add_df_table(doc, t2, font_size=6.5)

    doc.add_page_break()
    add_caption(doc, "Supplementary Table 3 | Visium breast section-held-out results.", "Section-held-out aggregate results for the two public Visium breast sections.", "source_data/SupplementaryTable3_VisiumBreastSectionHeldOutSummary.csv")
    add_df_table(doc, t3, font_size=7.5)

    doc.add_paragraph()
    add_caption(doc, "Supplementary Table 4 | Mixed-effects statistical summary.", "Mixed-effects models included Moran's I and model class as fixed effects and dataset as a random intercept.", "source_data/SupplementaryTable4_MixedEffectsSummary.csv")
    add_df_table(doc, t4, font_size=7.2)

    doc.add_paragraph()
    add_caption(doc, "Supplementary Table 5 | Boundary and non-interpretable cases.", "Cases retained as unavailable or non-interpretable rather than converted to zero.", "source_data/SupplementaryTable5_BoundaryCases.csv")
    add_df_table(doc, t5_dedup, font_size=6.7)

    doc.add_page_break()
    add_caption(doc, "Supplementary Table 6 | Two-sided paired Wilcoxon signed-rank results.", "Exact paired n values, Wilcoxon W statistics, P values and Benjamini-Hochberg false-discovery-rate-adjusted q values are shown for predefined random-versus-strict comparison families. The null hypothesis was median paired random-minus-strict difference = 0. Wilcoxon degrees of freedom are not applicable. Full effect columns are provided in the source-data CSV.", "source_data/SupplementaryTable6_WilcoxonSignedRankResults.csv")
    t6_display = t6[[
        "Dataset",
        "Comparison family",
        "Model",
        "Paired n",
        "Median random-minus-strict difference",
        "Wilcoxon W",
        "P value",
        "BH-FDR q value",
        "BH-FDR q < 0.05",
    ]]
    add_df_table(doc, t6_display, font_size=6.2)

    out = PORTAL / "Supplementary_Information_V4.docx"
    doc.save(out)
    return out


def update_source_index():
    idx = pd.read_csv(SD / "SourceData_Index.csv")
    additions = pd.DataFrame([
        ["Supplementary Figure 2", "all", "DLPFC; Visium breast", "PCA+Ridge; Spatial kNN", "random-size-matched versus spatial-buffer loss", "SupplementaryFigure2_SourceData.csv", "scripts/build_supplementary_v4.py", "PASS"],
        ["Supplementary Figure 3", "all", "DLPFC; Andersson; Thrane; Visium breast", "PCA+Ridge; Spatial kNN; GraphSAGE", "shared_panel_50 RLI and retention", "SupplementaryFigure3_SourceData.csv", "scripts/build_supplementary_v4.py", "PASS"],
        ["Supplementary Table 1", "all", "all", "all", "dataset/sample structure", "SupplementaryTable1_DatasetSampleStructure.csv", "scripts/build_supplementary_v4.py", "PASS"],
        ["Supplementary Table 2", "all", "spatial-buffer datasets", "all", "split sample counts", "SupplementaryTable2_SplitSampleCounts.csv", "scripts/build_supplementary_v4.py", "PASS"],
        ["Supplementary Table 3", "all", "Visium breast", "PCA+Ridge; Spatial kNN", "section-held-out summary", "SupplementaryTable3_VisiumBreastSectionHeldOutSummary.csv", "scripts/build_supplementary_v4.py", "PASS"],
        ["Supplementary Table 4", "all", "all", "all", "mixed-effects summary", "SupplementaryTable4_MixedEffectsSummary.csv", "scripts/build_supplementary_v4.py", "PASS"],
        ["Supplementary Table 5", "all", "all", "all", "boundary/non-interpretable cases", "SupplementaryTable5_BoundaryCases.csv", "scripts/build_supplementary_v4.py", "PASS"],
        ["Supplementary Table 6", "all", "DLPFC; Andersson; Thrane; Visium breast", "Mean; PCA+Ridge; Spatial kNN", "two-sided paired Wilcoxon signed-rank statistics with BH-FDR q values", "SupplementaryTable6_WilcoxonSignedRankResults.csv", "scripts/build_supplementary_v4.py", "PASS"],
    ], columns=idx.columns)
    idx = pd.concat([idx[~idx["figure"].isin(additions["figure"])], additions], ignore_index=True)
    idx.to_csv(SD / "SourceData_Index.csv", index=False)


def main():
    fig1, _ = build_supp_fig1()
    fig2, _ = build_supp_fig2()
    fig3, _ = build_supp_fig3()
    tables = make_tables()
    update_source_index()
    out = build_si_doc([fig1, fig2, fig3], tables)
    print(out)


if __name__ == "__main__":
    main()
