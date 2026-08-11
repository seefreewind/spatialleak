#!/usr/bin/env python3
"""Build a Word draft for the Nature Communications SpatialLeak manuscript."""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(".")
MANUSCRIPT = ROOT / "manuscript" / "SPATIALLEAK_NATCOMM_V6.md"
OUT = ROOT / "submission" / "nature_communications" / "SPATIALLEAK_NATCOMM_WORD_DRAFT.docx"
FIG_DIR = ROOT / "submission" / "nature_communications" / "FIGURES"


AUTHORS = [
    ("Yu Zhang", "1"),
    ("Ying Chen", "2"),
    ("Yue Liu", "2"),
    ("Da Lin", "1"),
]

AFFILIATIONS = [
    ("1", "Department of Ophthalmology, The Second Affiliated Hospital of Wenzhou Medical University, No. 109 Xueyuan West Road, Lucheng District, Wenzhou, Zhejiang Province, China"),
    ("2", "Wenzhou Medical University, Wenzhou, Zhejiang Province, China"),
]

CORRESPONDENCE = "Correspondence: Da Lin, 212574@wzhealth.com; ORCID 0009-0009-4410-0218"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="D9E2EC") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, widths):
    table.autofit = False
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def clean_inline(text: str) -> str:
    text = text.replace("`", "")
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"\[(.*?)\]\((.*?)\)", r"\1", text)
    return text


def split_markdown(md: str):
    lines = md.splitlines()
    sections = []
    current_title = None
    current = []
    for line in lines:
        if line.startswith("# "):
            continue
        if line.startswith("## "):
            if current_title is not None:
                sections.append((current_title, "\n".join(current).strip()))
            current_title = line[3:].strip()
            current = []
        else:
            current.append(line)
    if current_title is not None:
        sections.append((current_title, "\n".join(current).strip()))
    return sections


def add_title_block(doc: Document) -> None:
    title = "Evaluation design reshapes apparent generalization in spatial omics prediction"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(title)
    r.font.name = "Calibri"
    r.font.size = Pt(20)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string("0B2545")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    for i, (name, aff) in enumerate(AUTHORS):
        if i:
            p.add_run(", ")
        run = p.add_run(name)
        run.bold = True
        sup = p.add_run(aff)
        sup.font.superscript = True

    for idx, text in AFFILIATIONS:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        sup = p.add_run(idx)
        sup.font.superscript = True
        p.add_run(" " + text)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(CORRESPONDENCE)
    r.italic = True
    r.font.color.rgb = RGBColor.from_string("1F4D78")

    table = doc.add_table(rows=3, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    set_table_width(table, [1.85, 4.45])
    rows = [
        ("Article type", "Article; initial Word draft for Nature Communications submission preparation"),
        ("Conflict of interest", "The authors declare no competing interests."),
        ("Repository / DOI", "GitHub repository URL and Zenodo DOI to be inserted after public release."),
    ]
    for row, (k, v) in zip(table.rows, rows):
        set_cell_shading(row.cells[0], "F2F4F7")
        row.cells[0].paragraphs[0].add_run(k).bold = True
        row.cells[1].paragraphs[0].add_run(v)


def add_paragraph_text(doc: Document, text: str) -> None:
    text = clean_inline(text).strip()
    if not text:
        return
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Inches(0)
    p.add_run(text)


def add_figure(doc: Document, num: str, image: Path, caption: str) -> None:
    if not image.exists():
        add_paragraph_text(doc, f"Figure {num}. {caption} [Figure file pending final export.]")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image), width=Inches(6.2))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cap.paragraph_format.space_after = Pt(10)
    r = cap.add_run(f"Figure {num}. ")
    r.bold = True
    cap.add_run(caption)


def add_markdown_body(doc: Document, sections) -> None:
    skip = {"Author Contributions", "Funding", "Competing Interests", "Acknowledgements", "References"}
    for title, body in sections:
        if title in skip:
            continue
        doc.add_heading(title, level=1)
        blocks = [b.strip() for b in body.split("\n\n") if b.strip()]
        for block in blocks:
            if block.startswith("### "):
                doc.add_heading(clean_inline(block[4:].strip()), level=2)
                if "SpatialLeak first tested" in body and title == "Results":
                    pass
                continue
            if block.startswith("```"):
                continue
            if block.startswith("- "):
                for item in block.splitlines():
                    if item.startswith("- "):
                        doc.add_paragraph(clean_inline(item[2:].strip()), style="List Bullet")
                continue
            add_paragraph_text(doc, block)

            if title == "Results" and "SpatialLeak first tested whether random spot-level performance" in block:
                add_figure(
                    doc,
                    "1",
                    FIG_DIR / "Figure1_final.png",
                    "Evaluation design determines the generalization claim. The schematic separates permissive random spot evaluation, sources of apparent performance, isolation strategies and a six-tier evidence hierarchy.",
                )
            if title == "Results" and "Figure 3 summarizes the central heterogeneity result" in block:
                add_figure(
                    doc,
                    "3",
                    FIG_DIR / "Figure3_final_matrix.png",
                    "Two-channel landscape of apparent generalization inflation. Spatial-channel and patient-associated RLI are shown separately; NA denotes an unavailable or non-interpretable tier and is not treated as zero.",
                )


def add_back_matter(doc: Document) -> None:
    doc.add_heading("Author Contributions", level=1)
    contribs = [
        ("Yu Zhang", "conceived the study with Da Lin, designed the evaluation framework, performed the main analyses, prepared figures and drafted the manuscript."),
        ("Ying Chen", "contributed to data organization, preprocessing review, result checking and manuscript revision."),
        ("Yue Liu", "contributed to data curation, source-data preparation, reproducibility checks and manuscript revision."),
        ("Da Lin", "supervised the study, contributed to study design and interpretation, and reviewed and approved the manuscript for submission."),
    ]
    for name, text in contribs:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(name + ": ")
        r.bold = True
        p.add_run(text)

    doc.add_heading("Funding", level=1)
    add_paragraph_text(doc, "[Funding statement to be added.]")

    doc.add_heading("Competing Interests", level=1)
    add_paragraph_text(doc, "The authors declare no competing interests.")

    doc.add_heading("Acknowledgements", level=1)
    add_paragraph_text(doc, "[Acknowledgements to be added.]")

    doc.add_heading("References", level=1)
    add_paragraph_text(doc, "References are maintained in manuscript/references_master.bib and should be converted to final Nature style before submission.")


def main() -> None:
    md = read_md = MANUSCRIPT.read_text()
    sections = split_markdown(read_md)
    doc = Document()
    configure_styles(doc)
    add_title_block(doc)
    doc.add_section(WD_SECTION.NEW_PAGE)
    add_markdown_body(doc, sections)
    add_back_matter(doc)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
