#!/usr/bin/env python3
"""Build the V9 language-polished Nature Communications manuscript DOCX.

This script does not run analyses. It converts the locked Markdown manuscript
into a Word file and places the already generated main figures near first use.
"""
from __future__ import annotations

from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


SUB = Path("submission/nature_communications")
FIGS = SUB / "FIGURES"
MANUSCRIPT = SUB / "SpatialLeak_NatCommun_V9_language_polished.md"
OUT = SUB / "SpatialLeak_NatCommun_V9_language_polished.docx"

TITLE = "Evaluation design reshapes apparent generalization in spatial omics prediction"
AUTHORS = "Yu Zhang1, Ying Chen2, Yue Liu2, Da Lin1"
AFFILIATIONS = [
    "1 Department of Ophthalmology, The Second Affiliated Hospital of Wenzhou Medical University, No. 109 Xueyuan West Road, Lucheng District, Wenzhou, Zhejiang Province, China",
    "2 Wenzhou Medical University, Wenzhou, Zhejiang Province, China",
]
CORRESPONDENCE = "Correspondence: Da Lin, 212574@wzhealth.com; ORCID 0009-0009-4410-0218"


def add_figure(doc: Document, path: Path, caption: str) -> None:
    if path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run()
        r.add_picture(str(path), width=Inches(6.1))
    cp = doc.add_paragraph()
    cr = cp.add_run(caption)
    cr.italic = True
    cp.paragraph_format.space_after = Pt(8)


def build_docx() -> Path:
    text = MANUSCRIPT.read_text()
    doc = Document()
    props = doc.core_properties
    props.author = ""
    props.last_modified_by = ""
    props.title = TITLE
    props.subject = "SpatialLeak Nature Communications manuscript"
    props.keywords = "spatial omics; data leakage; evaluation design"
    props.comments = ""
    props.category = ""
    props.created = datetime(2026, 8, 11, tzinfo=timezone.utc)
    props.modified = datetime(2026, 8, 11, tzinfo=timezone.utc)

    sec = doc.sections[0]
    for margin in ["top_margin", "bottom_margin", "left_margin", "right_margin"]:
        setattr(sec, margin, Inches(1))

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    for name, size in [("Heading 1", 16), ("Heading 2", 13)]:
        styles[name].font.name = "Calibri"
        styles[name].font.size = Pt(size)
        styles[name].font.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(TITLE)
    r.bold = True
    r.font.size = Pt(16)
    for line in [AUTHORS, *AFFILIATIONS, CORRESPONDENCE]:
        pp = doc.add_paragraph(line)
        pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    for raw in text.splitlines():
        line = raw.strip()
        if not line or line.startswith("# "):
            continue
        if line in {AUTHORS, *AFFILIATIONS, CORRESPONDENCE}:
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:], level=1)
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:], level=2)
            continue

        doc.add_paragraph(line)

        if line.startswith("SpatialLeak first tested"):
            add_figure(
                doc,
                FIGS / "Figure1_final.png",
                "Figure 1. Evaluation design determines the generalization claim. (a) Random spot splitting can place neighboring training and test spots within the same local tissue context. (b) Observed test performance under permissive evaluation may combine local spatial dependence, subject-associated structure and transportable signal. (c) Increasing separation adds spatial buffers, section separation, subject separation (patient or donor), dataset separation and platform separation. (d) The evidence hierarchy links each evaluation tier to the level of generalization it can support; signal retained under stricter tiers provides stronger evidence for transportable signal.",
            )
            add_figure(
                doc,
                FIGS / "Figure2_final.png",
                "Figure 2. Predictive performance attenuates under stricter evaluation tiers. (a) Subject-associated evaluation compares random spot-level performance with patient- or donor-held-out performance in datasets supporting subject-level separation. (b) Spatial evaluation compares random performance with minimum graph distance ≥5 buffered spatial evaluation in datasets supporting within-section separation. Points indicate mean Pearson correlation across target genes, and connecting lines show the change between random and the corresponding stricter evaluation regime; Δr denotes random minus strict-tier Pearson correlation. Error bars indicate ±1 standard deviation (s.d.); random and spatial-buffer estimates summarize predefined seeds, whereas subject-held-out estimates summarize held-out subject groups as detailed in Source Data. Because dispersion units differ between random seed summaries and subject-held-out group summaries, error-bar widths should not be interpreted as directly comparable measures of sampling uncertainty. Dataset-model combinations with near-zero random performance, for which relative inflation is not interpretable, are excluded from the main display and reported in the Supplementary Information.",
            )
        if line.startswith("The subject-associated datasets"):
            add_figure(
                doc,
                FIGS / "Figure3_final_matrix.png",
                "Figure 3. SpatialLeak reveals heterogeneous channels of apparent generalization inflation across datasets and model classes. Rows represent interpretable dataset-model combinations, grouped by dataset. The spatial-neighborhood channel reports relative leakage inflation (RLI) between random and the prespecified hop-buffered spatial evaluation, prioritizing +5-hop where resolvable; exact strict splits are listed in Source Data. The subject-associated channel reports RLI between random and patient- or donor-held-out evaluation. Cell values show RLI, with stronger shading indicating larger positive evaluation-dependent attenuation. <0 denotes a negative RLI and therefore no positive inflation under the corresponding contrast. Hatched NA cells indicate evaluation tiers that were unavailable from the dataset structure or non-interpretable under the prespecified near-zero random-performance rule; NA values are not treated as zero. Descriptive pattern labels summarize the observed profile and do not represent threshold-based classifications.",
            )
        if line.startswith("SpatialLeak next tested"):
            add_figure(
                doc,
                FIGS / "Figure4_final.png",
                "Figure 4. Increasing spatial exclusion reveals dataset-dependent local neighborhood dependence. (a) Spatial kNN performance in DLPFC under random evaluation and increasingly buffered spatial splits. (b) Spatial kNN in dense Visium breast data, showing pronounced attenuation with increasing exclusion distance. (c) PCA+Ridge in the independent GSE278936 Visium cohort, where block-only separation produced little change relative to random evaluation, whereas non-zero graph-distance buffers reduced performance. Points show mean Pearson correlation across target genes; error bars indicate ±1 s.d. across frozen seeds (n = 10 for DLPFC and Visium breast; n = 5 for GSE278936). Random evaluation is shown as a permissive reference. Minimum graph distance ≥2 and ≥5 denote retained test spots whose nearest training spot is at least two or five graph edges away, respectively. Random-size-matched controls are reported in the Supplementary Information.",
            )

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build_docx())
